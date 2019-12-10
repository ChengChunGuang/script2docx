# -*- coding: utf-8 -*-

import re  # built-in
import logging # built-in
from docx import Document # pip install python-docx (note: python-docx depends on lxml)
from docx.oxml.ns import qn # python-docx
from docx.shared import Pt # python-docx

logging.basicConfig(level = logging.DEBUG,format = '[%(asctime)s] [%(levelname)5s] - %(message)s')
logger = logging.getLogger(__name__)

class sql_table():
    def __init__(self,table_name,table_comment,table_fields):
        self.table_name=table_name
        self.table_comment=table_comment
        self.table_fields=table_fields

def text_wrapper(text):
    if text: return text
    else : return r'-'

class sql_field():
    def __init__(self,field_name,filed_type,field_dflt,field_if_null,field_comment):
        self.field_name=field_name
        self.field_type=filed_type
        self.field_dflt=field_dflt
        self.field_if_null=field_if_null
        self.field_comment = field_comment

    def to_string(self):
        return r"[field name]: "+text_wrapper(self.field_name)+", [type]: "+text_wrapper(self.field_type)+ \
               ", [default value]: "+text_wrapper(self.field_dflt) + ", [is null]:" +text_wrapper(self.field_if_null)+ \
               ", [comment] "+text_wrapper(self.field_comment)

def find_and_clean(line,ptn):
    obj=ptn.search(line)
    if obj: return (re.sub(ptn,r' ',line),obj.group(1))
    else: return (line,None)

def parse_single_line(line):
    line = line.strip()
    if not line: return None

    field_pattern=re.compile(r'^\`(.+)\`') # field name
    (line, field_name)=find_and_clean(line,field_pattern)
    if not field_name: return None

    comment_pattern = re.compile(r'COMMENT\s+?\'(.+)\'') # comment
    (line, comment) = find_and_clean(line, comment_pattern)

    default_pattern=re.compile(r'DEFAULT\s+?(\S+)\s+') # default value
    (line, field_dflt) = find_and_clean(line, default_pattern)

    not_null_pattern=re.compile(r'(NOT\s+NULL)') # cheeck if the field can ben NULL
    (line, not_null) = find_and_clean(line, not_null_pattern)
    field_if_null=r'是'
    if not_null: field_if_null=r'否'

    field_type=re.sub(re.compile(r'\s|,'),r'',line.strip()) #field type
    return sql_field(field_name,field_type,field_dflt,field_if_null,comment)

def parse_single_table(table):
    (table_name,table_content,table_comment)=table
    logger.info("Table: "+table_name)
    lines = table_content.split('\n')
    table_fields=[]
    for line in lines:
        field=parse_single_line(line)
        if field:
            table_fields.append(field)
            logger.debug("  Field info: " + field.to_string())
    logger.info(str(len(table_fields))+" fields are found.")
    return sql_table(table_name,table_comment,table_fields)

def parse_sql_script(sql_script):
    fp = open(sql_script, 'r',encoding='utf-8')
    fc = fp.read()
    table_block_pattern = re.compile(
        r'^CREATE TABLE\s+?\`(.+?)\`\s+?\($(.*?)^\s*?\)\s+?ENGINE=InnoDB.*?COMMENT=\'(.+?)\'.*?;$',
        re.MULTILINE | re.DOTALL)
    result = table_block_pattern.findall(fc)
    logger.info(str(len(result))+" tables are found.")
    sql_tables = []
    for t in result:
        sql_tables.append(parse_single_table(t))
    return sql_tables

def dump_table_header(sql_table,table):
    table.cell(0,0).text=r'表名'
    table.cell(0,1).text = text_wrapper(sql_table.table_name)
    table.cell(1,0).text = r'功能'
    table.cell(1,1).text = text_wrapper(sql_table.table_comment)

def fill_table_row_content(table,row,row_content):
    col = 0
    while col < len(row_content):
        table.cell(row,col).text =row_content[col]
        col+=1

def dump_table_field_header(table,table_field_header):
    fill_table_row_content(table,2,table_field_header)

def dump_table_fields(sql_table,table):
    field_id=0
    for table_field in sql_table.table_fields:
        row_content=[text_wrapper(table_field.field_name),text_wrapper(table_field.field_type), \
                    text_wrapper(table_field.field_if_null),text_wrapper(table_field.field_dflt),\
                    text_wrapper(table_field.field_comment)]
        fill_table_row_content(table,3 + field_id,row_content)
        field_id+=1

def format_table(method):
    def formatter(*args, **kwargs):
        table=method(*args, **kwargs)
        table.style='Table Grid'
        table.autofit = True
        if len(table.columns) >= 4:
            table.columns[2].width = Pt(10)
            table.columns[3].width = Pt(25)
        table.cell(0, 1).merge(table.cell(0, len(table.columns) - 1))
        table.cell(1, 1).merge(table.cell(1, len(table.columns) - 1))
        return table
    return formatter

@format_table
def create_table(doc,row,col):
    return doc.add_table(rows=row, cols=col)

def dump_single_table(sql_table,doc):
    table_header=text_wrapper(sql_table.table_name) + '(' + text_wrapper(sql_table.table_comment) + ")"
    doc.add_heading(table_header)

    table_field_header = [r'字段', r'类型', r'Null', r'默认', r'注释']
    table=create_table(doc,len(sql_table.table_fields)+3,len(table_field_header))
    dump_table_header(sql_table,table) # the 1st and 2nd rows
    dump_table_field_header(table,table_field_header) # the 3rd row
    dump_table_fields(sql_table,table)
    doc.add_paragraph(r'')

def format_document(method):
    def formatter(*args, **kwargs):
        doc=method(*args, **kwargs)
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal'].font.size = Pt(10.5)
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        return doc
    return formatter

@format_document
def create_document():
    return Document()

def dump_to_file(sql_tables,dump_file):
    doc = create_document()
    logger.info("Beginning to dump sql tables to docx file.")
    for table in sql_tables:
        dump_single_table(table,doc)
    doc.save(dump_file)
    logger.info("Tables are dumped into " + dump_file)

if __name__ == "__main__":
    (sql_script_file,dump_to_docx)=(r'd:\table.sql',r'd:\dump.docx')
    dump_to_file(parse_sql_script(sql_script_file), dump_to_docx)

